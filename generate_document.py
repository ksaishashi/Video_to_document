from docx import Document
from docx.shared import Inches
import os

def create_document(output_path, descriptions, texts, diarization, aligned_speech, frames):
    doc = Document()
    doc.add_heading('RYBELSUS® Video Analysis Report', level=1)

    for frame_path, frame_name in zip(frames, sorted(descriptions.keys())):
        doc.add_picture(frame_path, width=Inches(2))  # Add frame image
        doc.add_paragraph(f"Description: {descriptions[frame_name]}", style='Normal')
        doc.add_paragraph(f"OCR Text: {texts[frame_name]}", style='IntenseQuote')

        # Add aligned speech text
        if frame_name in aligned_speech:
            speaker_text = aligned_speech[frame_name].strip()
            doc.add_paragraph(f"Speech: {speaker_text}", style='Quote')

    doc.save(output_path)
